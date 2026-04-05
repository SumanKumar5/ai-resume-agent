from docx import Document
from docx.shared import Pt
import os


def read_resume(file_path):
    """
    Reads a .docx resume and returns full text.
    """
    try:
        doc = Document(file_path)

        full_text = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)

        return "\n".join(full_text)

    except Exception as e:
        print(f"[ERROR] Failed to read resume: {e}")
        return ""


def save_resume(text, job, output_dir="output/resumes"):
    """
    Saves tailored resume with improved formatting.
    """
    try:
        os.makedirs(output_dir, exist_ok=True)

        # Clean filename
        title = job["title"].replace(" ", "_").replace("/", "").lower()
        file_path = os.path.join(output_dir, f"resume_{title}.docx")

        doc = Document()

        for line in text.split("\n"):
            line = line.strip()

            if not line:
                continue

            # Remove markdown symbols
            clean_line = line.replace("*", "")

            # Detect headings
            if line.isupper() or line.startswith("**"):
                para = doc.add_paragraph()
                run = para.add_run(clean_line)
                run.bold = True
                run.font.size = Pt(14)

            # Detect sub-headings (like job titles)
            elif "|" in line or line.startswith("•"):
                para = doc.add_paragraph(clean_line)
                para.runs[0].bold = True

            else:
                doc.add_paragraph(clean_line)

        doc.save(file_path)

        print(f"[SUCCESS] Resume saved: {file_path}")

        return file_path

    except Exception as e:
        print(f"[ERROR] Failed to save resume: {e}")
        return None