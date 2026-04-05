import json
import logging
import openpyxl
from docx import Document
from pathlib import Path

logger = logging.getLogger(__name__)


def load_excel(filepath: str) -> list[dict]:
    """Read job links Excel file and return list of job link records."""
    try:
        wb = openpyxl.load_workbook(filepath)
        ws = wb.active
        headers = [cell.value for cell in ws[1]]
        records = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            record = dict(zip(headers, row))
            records.append(record)
        logger.info(f"Loaded {len(records)} records from Excel.")
        return records
    except Exception as e:
        logger.error(f"Failed to load Excel file: {e}")
        raise


def load_json(filepath: str) -> list[dict]:
    """Read jobs JSON file and return list of job detail records."""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            data = json.load(f)
        jobs = data.get("jobs", [])
        logger.info(f"Loaded {len(jobs)} jobs from JSON.")
        return jobs
    except Exception as e:
        logger.error(f"Failed to load JSON file: {e}")
        raise


def merge_jobs(excel_records: list[dict], json_jobs: list[dict]) -> list[dict]:
    """Join Excel and JSON records on # / id field into one combined record per job."""
    json_map = {job["id"]: job for job in json_jobs}
    merged = []
    for record in excel_records:
        job_id = record.get("#")
        if job_id in json_map:
            combined = {**record, **json_map[job_id]}
            merged.append(combined)
        else:
            logger.warning(f"No JSON match found for job id: {job_id}")
    logger.info(f"Merged {len(merged)} complete job records.")
    return merged


def load_resume(filepath: str) -> str:
    """Extract full text content from the candidate resume .docx file."""
    try:
        doc = Document(filepath)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
        resume_text = "\n".join(full_text)
        logger.info("Resume loaded successfully.")
        return resume_text
    except Exception as e:
        logger.error(f"Failed to load resume: {e}")
        raise


def load_all_inputs(
    excel_path: str,
    json_path: str,
    resume_path: str
) -> tuple[list[dict], str]:
    """
    Master function: loads and merges all inputs.
    Returns (merged_jobs, resume_text).
    """
    excel_records = load_excel(excel_path)
    json_jobs = load_json(json_path)
    merged_jobs = merge_jobs(excel_records, json_jobs)
    resume_text = load_resume(resume_path)
    return merged_jobs, resume_text