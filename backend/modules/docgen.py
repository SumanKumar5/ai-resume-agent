import logging
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

logger = logging.getLogger(__name__)


def set_font(run, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_section_heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text.upper())
    set_font(run, bold=True, size=11, color=(0, 0, 0))
    border = para._element
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = border.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def add_bullet(doc, text, indent=0.3):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Inches(indent)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    set_font(run, size=10)
    return para


def build_docx(data: dict, output_path: str):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(data["name"])
    set_font(name_run, bold=True, size=16)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_before = Pt(2)
    contact_run = contact_para.add_run(data["contact"])
    set_font(contact_run, size=9, color=(80, 80, 80))

    add_section_heading(doc, "Professional Summary")
    summary_para = doc.add_paragraph()
    summary_para.paragraph_format.space_after = Pt(4)
    summary_run = summary_para.add_run(data["summary"])
    set_font(summary_run, size=10)

    add_section_heading(doc, "Technical Skills")
    for category, values in data["skills"].items():
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        label_run = para.add_run(f"{category}: ")
        set_font(label_run, bold=True, size=10)
        value_run = para.add_run(values)
        set_font(value_run, size=10)

    add_section_heading(doc, "Work Experience")
    for exp in data["experience"]:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(1)
        title_run = para.add_run(f"{exp['title']}  |  {exp['company']}")
        set_font(title_run, bold=True, size=10)
        period_para = doc.add_paragraph()
        period_para.paragraph_format.space_after = Pt(2)
        period_run = period_para.add_run(exp["period"])
        set_font(period_run, size=9, color=(100, 100, 100))
        period_run.italic = True
        for bullet in exp["bullets"]:
            add_bullet(doc, bullet)

    add_section_heading(doc, "Projects")
    for project in data["projects"]:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(1)
        proj_run = para.add_run(f"{project['name']}  —  ")
        set_font(proj_run, bold=True, size=10)
        url_run = para.add_run(project["url"])
        set_font(url_run, size=10, color=(0, 70, 180))
        for bullet in project["bullets"]:
            add_bullet(doc, bullet)

    add_section_heading(doc, "Education")
    edu = data["education"]
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(1)
    edu_run = para.add_run(f"{edu['university']}  —  {edu['degree']}")
    set_font(edu_run, bold=True, size=10)
    period_para = doc.add_paragraph()
    period_para.paragraph_format.space_after = Pt(2)
    period_run = period_para.add_run(f"{edu['period']}  |  GPA: {edu['gpa']}")
    set_font(period_run, size=9, color=(100, 100, 100))
    period_run.italic = True
    for detail in edu["details"]:
        add_bullet(doc, detail)

    if data.get("certifications"):
        add_section_heading(doc, "Certifications & Activities")
        for cert in data["certifications"]:
            add_bullet(doc, cert)

    doc.save(output_path)
    logger.info(f"DOCX saved: {output_path}")


def generate_resume(data: dict, job_title: str, candidate_name: str, output_dir: str) -> str:
    try:
        from datetime import datetime
        Path(output_dir).mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = job_title.lower().replace(" ", "_").replace("/", "_")
        safe_name = candidate_name.lower().replace(" ", "_").replace(".", "")
        base_filename = f"resume_{safe_name}_{safe_title}_{timestamp}"

        docx_path = os.path.join(output_dir, f"{base_filename}.docx")
        pdf_path = os.path.join(output_dir, f"{base_filename}.pdf")

        build_docx(data, docx_path)
        convert(docx_path, pdf_path)
        logger.info(f"PDF saved: {pdf_path}")

        return pdf_path

    except Exception as e:
        logger.error(f"Document generation failed for {job_title}: {e}")
        raise